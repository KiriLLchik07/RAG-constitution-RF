import docx
from collections import defaultdict

doc = docx.Document("data/raw/constitutionrf.docx")

structure = defaultdict(list)
current_chapter = "Преамбула"
current_article = None

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    
    if text.startswith("ГЛАВА "):
        current_chapter = text
        print(f"Обнаружена глава: {current_chapter}")
    
    elif text.startswith("Статья "):
        current_article = text.split(".")[0]
        structure[current_chapter].append(current_article)
        print(f"Статья: {current_article} в {current_chapter}")

print("\nСтатистика:")
for chapter, articles in structure.items():
    print(f"{chapter}: {len(articles)} статей")

clean_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
with open("data/raw/constitution_rf_clean.txt", "w", encoding="utf-8") as f:
    f.write(clean_text)
